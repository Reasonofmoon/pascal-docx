"""
파스칼 DOCX 교재 생성 모듈
CSV 데이터를 입력받아 출판 수준의 완전한 DOCX 교재를 생성하는 모듈
"""

import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from datetime import datetime
import json
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
from enum import Enum
import logging

# 로깅 설정
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class EducationLevel(str, Enum):
    """교육 레벨 정의"""
    PREPARATION = "preparation"
    REGULAR = "regular"
    MASTERY = "mastery"

@dataclass
class ColorTheme:
    """색상 테마 정의"""
    primary: str
    secondary: str
    accent: str
    text: str
    background: str
    
    @classmethod
    def get_theme(cls, level: EducationLevel) -> 'ColorTheme':
        """레벨별 색상 테마 반환"""
        themes = {
            EducationLevel.PREPARATION: cls(
                primary="#4A90E2",      # 밝은 파란색
                secondary="#7ED321",    # 연두색
                accent="#F5A623",       # 주황색
                text="#333333",         # 진한 회색
                background="#F8F9FA"    # 연한 회색
            ),
            EducationLevel.REGULAR: cls(
                primary="#5856D6",      # 보라색
                secondary="#34C759",    # 초록색
                accent="#FF9500",       # 주황색
                text="#1D1D1F",         # 검은색
                background="#FFFFFF"    # 흰색
            ),
            EducationLevel.MASTERY: cls(
                primary="#1D1D1F",      # 검은색
                secondary="#8E8E93",    # 회색
                accent="#FF3B30",       # 빨간색
                text="#000000",         # 검은색
                background="#F2F2F7"    # 연한 회색
            )
        }
        return themes[level]

@dataclass
class DocumentSettings:
    """문서 설정"""
    title: str
    subtitle: str
    author: str
    institution: str
    level: EducationLevel
    book_title: str
    book_author: str
    ar_level: float
    creation_date: datetime
    
class DOCXGenerator:
    """DOCX 교재 생성 클래스"""
    
    def __init__(self):
        self.document = None
        self.color_theme = None
        self.settings = None
        
    def generate_textbook(
        self, 
        csv_path: str, 
        output_path: str, 
        settings: DocumentSettings
    ) -> str:
        """CSV 데이터로부터 완전한 교재 생성"""
        
        logger.info(f"Starting textbook generation from {csv_path}")
        
        # 1. CSV 데이터 로드 및 검증
        df = self._load_and_validate_csv(csv_path)
        
        # 2. 문서 초기화
        self._initialize_document(settings)
        
        # 3. 표지 생성
        self._create_cover_page()
        
        # 4. 목차 생성
        self._create_table_of_contents(df)
        
        # 5. 서문 생성
        self._create_preface()
        
        # 6. 각 토론 주제별 챕터 생성
        self._create_chapters(df)
        
        # 7. 부록 생성
        self._create_appendix(df)
        
        # 8. 평가 기준 생성
        self._create_assessment_criteria()
        
        # 9. 문서 저장
        self.document.save(output_path)
        
        logger.info(f"Textbook generated successfully: {output_path}")
        return output_path
    
    def _load_and_validate_csv(self, csv_path: str) -> pd.DataFrame:
        """CSV 파일 로드 및 검증"""
        try:
            df = pd.read_csv(csv_path, encoding='utf-8')
            
            # 필수 컬럼 확인
            required_columns = [
                'Topic_ID', 'Title', 'Description', 'Level', 'Area', 'Format',
                'Pro_Arguments', 'Con_Arguments', 'Background', 'Vocabulary', 'Time_Minutes'
            ]
            
            missing_columns = [col for col in required_columns if col not in df.columns]
            if missing_columns:
                raise ValueError(f"Missing required columns: {missing_columns}")
            
            logger.info(f"CSV loaded successfully with {len(df)} topics")
            return df
            
        except Exception as e:
            logger.error(f"Failed to load CSV: {e}")
            raise
    
    def _initialize_document(self, settings: DocumentSettings):
        """문서 초기화"""
        self.document = Document()
        self.settings = settings
        self.color_theme = ColorTheme.get_theme(settings.level)
        
        # 문서 속성 설정
        self.document.core_properties.title = settings.title
        self.document.core_properties.author = settings.author
        self.document.core_properties.subject = f"Pascal Debate Textbook - {settings.level.value.title()}"
        self.document.core_properties.created = settings.creation_date
        
        # 기본 스타일 설정
        self._setup_styles()
        
        # 페이지 설정
        self._setup_page_layout()
    
    def _setup_styles(self):
        """문서 스타일 설정"""
        styles = self.document.styles
        
        # 제목 스타일
        if 'Custom Title' not in [style.name for style in styles]:
            title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Arial'
            title_font.size = Pt(24)
            title_font.bold = True
            title_font.color.rgb = RGBColor.from_string(self.color_theme.primary.replace('#', ''))
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
        
        # 부제목 스타일
        if 'Custom Subtitle' not in [style.name for style in styles]:
            subtitle_style = styles.add_style('Custom Subtitle', WD_STYLE_TYPE.PARAGRAPH)
            subtitle_font = subtitle_style.font
            subtitle_font.name = 'Arial'
            subtitle_font.size = Pt(18)
            subtitle_font.bold = True
            subtitle_font.color.rgb = RGBColor.from_string(self.color_theme.secondary.replace('#', ''))
            subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_style.paragraph_format.space_after = Pt(10)
        
        # 챕터 제목 스타일
        if 'Chapter Title' not in [style.name for style in styles]:
            chapter_style = styles.add_style('Chapter Title', WD_STYLE_TYPE.PARAGRAPH)
            chapter_font = chapter_style.font
            chapter_font.name = 'Arial'
            chapter_font.size = Pt(20)
            chapter_font.bold = True
            chapter_font.color.rgb = RGBColor.from_string(self.color_theme.primary.replace('#', ''))
            chapter_style.paragraph_format.space_before = Pt(24)
            chapter_style.paragraph_format.space_after = Pt(12)
        
        # 섹션 제목 스타일
        if 'Section Title' not in [style.name for style in styles]:
            section_style = styles.add_style('Section Title', WD_STYLE_TYPE.PARAGRAPH)
            section_font = section_style.font
            section_font.name = 'Arial'
            section_font.size = Pt(16)
            section_font.bold = True
            section_font.color.rgb = RGBColor.from_string(self.color_theme.accent.replace('#', ''))
            section_style.paragraph_format.space_before = Pt(12)
            section_style.paragraph_format.space_after = Pt(6)
        
        # 본문 스타일
        if 'Custom Body' not in [style.name for style in styles]:
            body_style = styles.add_style('Custom Body', WD_STYLE_TYPE.PARAGRAPH)
            body_font = body_style.font
            body_font.name = 'Arial'
            body_font.size = Pt(11)
            body_font.color.rgb = RGBColor.from_string(self.color_theme.text.replace('#', ''))
            body_style.paragraph_format.space_after = Pt(6)
            body_style.paragraph_format.line_spacing = 1.15
    
    def _setup_page_layout(self):
        """페이지 레이아웃 설정"""
        section = self.document.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
    
    def _create_cover_page(self):
        """표지 생성"""
        # 레벨별 표지 디자인
        level_designs = {
            EducationLevel.PREPARATION: "🌱 기초 단계",
            EducationLevel.REGULAR: "🌿 발전 단계", 
            EducationLevel.MASTERY: "🌳 숙달 단계"
        }
        
        # 메인 제목
        title_para = self.document.add_paragraph()
        title_para.style = 'Custom Title'
        title_run = title_para.runs[0] if title_para.runs else title_para.add_run()
        title_run.text = self.settings.title
        
        # 부제목
        subtitle_para = self.document.add_paragraph()
        subtitle_para.style = 'Custom Subtitle'
        subtitle_run = subtitle_para.runs[0] if subtitle_para.runs else subtitle_para.add_run()
        subtitle_run.text = self.settings.subtitle
        
        # 레벨 표시
        level_para = self.document.add_paragraph()
        level_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        level_run = level_para.add_run(level_designs[self.settings.level])
        level_run.font.size = Pt(16)
        level_run.font.bold = True
        level_run.font.color.rgb = RGBColor.from_string(self.color_theme.accent.replace('#', ''))
        
        # 원서 정보
        self.document.add_paragraph()  # 공백
        book_para = self.document.add_paragraph()
        book_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        book_run = book_para.add_run(f'Based on "{self.settings.book_title}" by {self.settings.book_author}')
        book_run.font.size = Pt(14)
        book_run.font.italic = True
        
        ar_para = self.document.add_paragraph()
        ar_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar_run = ar_para.add_run(f'AR Level: {self.settings.ar_level}')
        ar_run.font.size = Pt(12)
        
        # 기관 정보
        self.document.add_paragraph()  # 공백
        self.document.add_paragraph()  # 공백
        
        institution_para = self.document.add_paragraph()
        institution_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        institution_run = institution_para.add_run(self.settings.institution)
        institution_run.font.size = Pt(14)
        institution_run.font.bold = True
        
        author_para = self.document.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_para.add_run(f'작성자: {self.settings.author}')
        author_run.font.size = Pt(12)
        
        date_para = self.document.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(self.settings.creation_date.strftime('%Y년 %m월 %d일'))
        date_run.font.size = Pt(12)
        
        # 페이지 나누기
        self.document.add_page_break()
    
    def _create_table_of_contents(self, df: pd.DataFrame):
        """목차 생성"""
        toc_title = self.document.add_paragraph()
        toc_title.style = 'Chapter Title'
        toc_title.add_run('목차')
        
        # 서문
        toc_item = self.document.add_paragraph('서문', style='Custom Body')
        toc_item.paragraph_format.left_indent = Inches(0.25)
        
        # 각 챕터
        for idx, row in df.iterrows():
            chapter_num = idx + 1
            toc_item = self.document.add_paragraph(
                f'Chapter {chapter_num}: {row["Title"]}', 
                style='Custom Body'
            )
            toc_item.paragraph_format.left_indent = Inches(0.25)
        
        # 부록
        toc_item = self.document.add_paragraph('부록 A: 전체 어휘 목록', style='Custom Body')
        toc_item.paragraph_format.left_indent = Inches(0.25)
        
        toc_item = self.document.add_paragraph('부록 B: 평가 기준', style='Custom Body')
        toc_item.paragraph_format.left_indent = Inches(0.25)
        
        self.document.add_page_break()
    
    def _create_preface(self):
        """서문 생성"""
        preface_title = self.document.add_paragraph()
        preface_title.style = 'Chapter Title'
        preface_title.add_run('서문')
        
        # 레벨별 서문 내용
        level_prefaces = {
            EducationLevel.PREPARATION: """
            이 교재는 영어 원서를 통한 토론 학습의 첫 걸음을 위해 제작되었습니다. 
            기초 단계 학습자들이 영어로 자신의 생각을 표현하고 다른 사람의 의견을 듣는 
            경험을 통해 영어 실력과 사고력을 동시에 기를 수 있도록 구성되었습니다.
            
            각 챕터는 명확한 구조와 충분한 지원을 제공하여 학습자가 안전하고 
            자신감 있게 토론에 참여할 수 있도록 돕습니다.
            """,
            EducationLevel.REGULAR: """
            이 교재는 영어 토론 실력을 한 단계 발전시키고자 하는 학습자들을 위해 
            제작되었습니다. 기본적인 토론 기술을 바탕으로 더 깊이 있는 분석과 
            비판적 사고를 기를 수 있도록 구성되었습니다.
            
            다양한 관점을 고려하고 논리적으로 자신의 주장을 펼치는 능력을 
            기르는 것이 이 교재의 주요 목표입니다.
            """,
            EducationLevel.MASTERY: """
            이 교재는 고급 수준의 영어 토론 능력을 완성하고자 하는 학습자들을 위해 
            제작되었습니다. 복잡한 주제에 대한 독립적인 분석과 창의적인 사고를 
            통해 세계적 수준의 토론 실력을 기를 수 있도록 구성되었습니다.
            
            학습자는 이 교재를 통해 글로벌 무대에서 당당히 소통할 수 있는 
            역량을 갖추게 될 것입니다.
            """
        }
        
        preface_content = level_prefaces[self.settings.level]
        for paragraph in preface_content.strip().split('\n\n'):
            if paragraph.strip():
                p = self.document.add_paragraph(paragraph.strip(), style='Custom Body')
                p.paragraph_format.first_line_indent = Inches(0.25)
        
        # 사용법 안내
        usage_title = self.document.add_paragraph()
        usage_title.style = 'Section Title'
        usage_title.add_run('교재 사용법')
        
        usage_steps = [
            "1. 각 챕터의 배경 정보를 먼저 읽어보세요.",
            "2. 핵심 어휘를 학습하고 예문을 확인하세요.",
            "3. 토론 주제에 대해 찬반 양쪽 입장을 고려해보세요.",
            "4. 동료들과 함께 토론을 진행하세요.",
            "5. 토론 후 성찰 활동을 통해 학습을 정리하세요."
        ]
        
        for step in usage_steps:
            step_para = self.document.add_paragraph(step, style='Custom Body')
            step_para.paragraph_format.left_indent = Inches(0.25)
        
        self.document.add_page_break()
    
    def _create_chapters(self, df: pd.DataFrame):
        """각 토론 주제별 챕터 생성"""
        for idx, row in df.iterrows():
            chapter_num = idx + 1
            self._create_single_chapter(chapter_num, row)
            
            # 마지막 챕터가 아니면 페이지 나누기
            if idx < len(df) - 1:
                self.document.add_page_break()
    
    def _create_single_chapter(self, chapter_num: int, row: pd.Series):
        """단일 챕터 생성"""
        # 챕터 제목
        chapter_title = self.document.add_paragraph()
        chapter_title.style = 'Chapter Title'
        chapter_title.add_run(f'Chapter {chapter_num}: {row["Title"]}')
        
        # 기본 정보 테이블
        self._add_chapter_info_table(row)
        
        # 배경 정보
        self._add_background_section(row)
        
        # 핵심 어휘
        self._add_vocabulary_section(row)
        
        # 토론 준비
        self._add_debate_preparation_section(row)
        
        # 토론 진행
        self._add_debate_process_section(row)
        
        # 글쓰기 활동
        self._add_writing_activity_section(row)
        
        # 성찰 활동
        self._add_reflection_section(row)
    
    def _add_chapter_info_table(self, row: pd.Series):
        """챕터 정보 테이블 추가"""
        table = self.document.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 테이블 내용
        info_data = [
            ('영역', row['Area']),
            ('형식', row['Format'].replace('_', ' ').title()),
            ('난이도', row['Level'].title()),
            ('예상 시간', f"{row['Time_Minutes']}분"),
            ('주제 ID', row['Topic_ID'])
        ]
        
        for i, (label, value) in enumerate(info_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = str(value)
            
            # 헤더 셀 스타일링
            header_cell = table.cell(i, 0)
            header_para = header_cell.paragraphs[0]
            header_run = header_para.runs[0] if header_para.runs else header_para.add_run(label)
            header_run.font.bold = True
        
        self.document.add_paragraph()  # 공백
    
    def _add_background_section(self, row: pd.Series):
        """배경 정보 섹션 추가"""
        bg_title = self.document.add_paragraph()
        bg_title.style = 'Section Title'
        bg_title.add_run('📚 배경 정보')
        
        # 토론 주제 설명
        desc_para = self.document.add_paragraph(row['Description'], style='Custom Body')
        desc_para.paragraph_format.first_line_indent = Inches(0.25)
        
        # 배경 정보
        if pd.notna(row['Background']) and row['Background'].strip():
            bg_para = self.document.add_paragraph(row['Background'], style='Custom Body')
            bg_para.paragraph_format.first_line_indent = Inches(0.25)
    
    def _add_vocabulary_section(self, row: pd.Series):
        """핵심 어휘 섹션 추가"""
        vocab_title = self.document.add_paragraph()
        vocab_title.style = 'Section Title'
        vocab_title.add_run('📝 핵심 어휘')
        
        if pd.notna(row['Vocabulary']) and row['Vocabulary'].strip():
            vocab_list = [word.strip() for word in row['Vocabulary'].split('|')]
            
            # 어휘를 2열 테이블로 구성
            vocab_table = self.document.add_table(rows=(len(vocab_list) + 1) // 2, cols=2)
            vocab_table.style = 'Table Grid'
            
            for i, word in enumerate(vocab_list):
                row_idx = i // 2
                col_idx = i % 2
                cell = vocab_table.cell(row_idx, col_idx)
                cell.text = word
                
                # 어휘 스타일링
                cell_para = cell.paragraphs[0]
                cell_run = cell_para.runs[0] if cell_para.runs else cell_para.add_run(word)
                cell_run.font.bold = True
        
        self.document.add_paragraph()  # 공백
    
    def _add_debate_preparation_section(self, row: pd.Series):
        """토론 준비 섹션 추가"""
        prep_title = self.document.add_paragraph()
        prep_title.style = 'Section Title'
        prep_title.add_run('🤔 토론 준비')
        
        # 찬성 논거
        if pd.notna(row['Pro_Arguments']) and row['Pro_Arguments'].strip():
            pro_subtitle = self.document.add_paragraph('찬성 논거:', style='Custom Body')
            pro_subtitle.runs[0].font.bold = True
            pro_subtitle.runs[0].font.color.rgb = RGBColor.from_string(self.color_theme.secondary.replace('#', ''))
            
            pro_args = [arg.strip() for arg in row['Pro_Arguments'].split('|')]
            for i, arg in enumerate(pro_args, 1):
                arg_para = self.document.add_paragraph(f'{i}. {arg}', style='Custom Body')
                arg_para.paragraph_format.left_indent = Inches(0.5)
        
        # 반대 논거
        if pd.notna(row['Con_Arguments']) and row['Con_Arguments'].strip():
            con_subtitle = self.document.add_paragraph('반대 논거:', style='Custom Body')
            con_subtitle.runs[0].font.bold = True
            con_subtitle.runs[0].font.color.rgb = RGBColor.from_string(self.color_theme.accent.replace('#', ''))
            
            con_args = [arg.strip() for arg in row['Con_Arguments'].split('|')]
            for i, arg in enumerate(con_args, 1):
                arg_para = self.document.add_paragraph(f'{i}. {arg}', style='Custom Body')
                arg_para.paragraph_format.left_indent = Inches(0.5)
    
    def _add_debate_process_section(self, row: pd.Series):
        """토론 진행 섹션 추가"""
        process_title = self.document.add_paragraph()
        process_title.style = 'Section Title'
        process_title.add_run('💬 토론 진행')
        
        # 레벨별 토론 가이드
        level_guides = {
            EducationLevel.PREPARATION: [
                "1. 팀을 나누어 찬성팀과 반대팀을 구성합니다.",
                "2. 각 팀은 5분간 논거를 정리합니다.",
                "3. 찬성팀이 먼저 2분간 주장을 발표합니다.",
                "4. 반대팀이 2분간 주장을 발표합니다.",
                "5. 각 팀이 1분씩 반박 기회를 가집니다.",
                "6. 마지막으로 각 팀이 1분씩 최종 발언을 합니다."
            ],
            EducationLevel.REGULAR: [
                "1. 팀을 구성하고 역할을 분담합니다.",
                "2. 각 팀은 7분간 전략을 수립합니다.",
                "3. 개회사와 주제 소개 (2분)",
                "4. 찬성팀 주장 발표 (3분)",
                "5. 반대팀 주장 발표 (3분)",
                "6. 교차 질의 시간 (각 팀 2분씩)",
                "7. 반박 및 재반박 (각 팀 2분씩)",
                "8. 최종 발언 (각 팀 1분씩)"
            ],
            EducationLevel.MASTERY: [
                "1. 독립적인 연구와 준비 시간 (10분)",
                "2. 자유로운 토론 형식으로 진행",
                "3. 사회자가 토론을 조율하며 진행",
                "4. 논리적 근거와 반박에 중점",
                "5. 창의적 해결방안 모색",
                "6. 종합적 결론 도출"
            ]
        }
        
        guide_steps = level_guides[self.settings.level]
        for step in guide_steps:
            step_para = self.document.add_paragraph(step, style='Custom Body')
            step_para.paragraph_format.left_indent = Inches(0.25)
    
    def _add_writing_activity_section(self, row: pd.Series):
        """글쓰기 활동 섹션 추가"""
        writing_title = self.document.add_paragraph()
        writing_title.style = 'Section Title'
        writing_title.add_run('✍️ 글쓰기 활동')
        
        # 레벨별 글쓰기 가이드
        level_writing = {
            EducationLevel.PREPARATION: {
                "word_count": "150-200단어",
                "time": "30분",
                "structure": [
                    "도입: 주제 소개와 자신의 입장 (2-3문장)",
                    "본론: 두 가지 근거와 설명 (각 3-4문장)",
                    "결론: 입장 재확인과 마무리 (2-3문장)"
                ]
            },
            EducationLevel.REGULAR: {
                "word_count": "250-300단어",
                "time": "40분",
                "structure": [
                    "도입: 주제의 중요성과 논제 제시",
                    "본론 1: 첫 번째 근거와 구체적 예시",
                    "본론 2: 두 번째 근거와 반대 의견 고려",
                    "결론: 논거 요약과 함의 제시"
                ]
            },
            EducationLevel.MASTERY: {
                "word_count": "400-500단어",
                "time": "50분",
                "structure": [
                    "도입: 문제 제기와 논제의 복잡성 인식",
                    "본론: 다각적 분석과 비판적 평가",
                    "반박: 상대방 입장 고려와 재반박",
                    "결론: 종합적 판단과 미래 전망"
                ]
            }
        }
        
        writing_info = level_writing[self.settings.level]
        
        # 글쓰기 요구사항
        req_para = self.document.add_paragraph(
            f"분량: {writing_info['word_count']} | 시간: {writing_info['time']}", 
            style='Custom Body'
        )
        req_para.runs[0].font.bold = True
        
        # 구조 가이드
        structure_para = self.document.add_paragraph('글 구조:', style='Custom Body')
        structure_para.runs[0].font.bold = True
        
        for item in writing_info['structure']:
            item_para = self.document.add_paragraph(f'• {item}', style='Custom Body')
            item_para.paragraph_format.left_indent = Inches(0.5)
    
    def _add_reflection_section(self, row: pd.Series):
        """성찰 활동 섹션 추가"""
        reflection_title = self.document.add_paragraph()
        reflection_title.style = 'Section Title'
        reflection_title.add_run('🤗 성찰 활동')
        
        reflection_questions = [
            "이번 토론에서 가장 인상 깊었던 논거는 무엇인가요?",
            "상대방의 의견 중 고려해볼 만한 점이 있었나요?",
            "자신의 주장을 더 효과적으로 전달하려면 어떻게 해야 할까요?",
            "이 주제와 관련하여 더 알아보고 싶은 것이 있나요?",
            "오늘 학습한 어휘 중 가장 유용하다고 생각하는 단어는 무엇인가요?"
        ]
        
        for i, question in enumerate(reflection_questions, 1):
            q_para = self.document.add_paragraph(f'{i}. {question}', style='Custom Body')
            q_para.paragraph_format.left_indent = Inches(0.25)
            
            # 답변 공간
            answer_para = self.document.add_paragraph('답변: ___________________________', style='Custom Body')
            answer_para.paragraph_format.left_indent = Inches(0.5)
            answer_para.paragraph_format.space_after = Pt(12)
    
    def _create_appendix(self, df: pd.DataFrame):
        """부록 생성"""
        self.document.add_page_break()
        
        # 부록 A: 전체 어휘 목록
        appendix_title = self.document.add_paragraph()
        appendix_title.style = 'Chapter Title'
        appendix_title.add_run('부록 A: 전체 어휘 목록')
        
        # 모든 어휘 수집
        all_vocab = set()
        for _, row in df.iterrows():
            if pd.notna(row['Vocabulary']) and row['Vocabulary'].strip():
                vocab_list = [word.strip() for word in row['Vocabulary'].split('|')]
                all_vocab.update(vocab_list)
        
        # 알파벳 순으로 정렬
        sorted_vocab = sorted(list(all_vocab))
        
        # 3열 테이블로 어휘 목록 생성
        vocab_table = self.document.add_table(rows=(len(sorted_vocab) + 2) // 3, cols=3)
        vocab_table.style = 'Table Grid'
        
        for i, word in enumerate(sorted_vocab):
            row_idx = i // 3
            col_idx = i % 3
            cell = vocab_table.cell(row_idx, col_idx)
            cell.text = word
    
    def _create_assessment_criteria(self):
        """평가 기준 생성"""
        self.document.add_page_break()
        
        assessment_title = self.document.add_paragraph()
        assessment_title.style = 'Chapter Title'
        assessment_title.add_run('부록 B: 평가 기준')
        
        # 평가 영역
        criteria_areas = [
            {
                "name": "내용 이해도",
                "description": "주제에 대한 이해와 관련 지식 활용",
                "excellent": "주제를 완전히 이해하고 풍부한 배경지식 활용",
                "good": "주제를 잘 이해하고 적절한 지식 활용",
                "satisfactory": "주제를 기본적으로 이해하고 일부 지식 활용",
                "needs_improvement": "주제 이해가 부족하고 지식 활용 미흡"
            },
            {
                "name": "논리적 사고",
                "description": "논거의 논리성과 근거의 적절성",
                "excellent": "논리적이고 설득력 있는 논거 제시",
                "good": "대체로 논리적인 논거 제시",
                "satisfactory": "기본적인 논거 제시",
                "needs_improvement": "논거가 불명확하거나 논리성 부족"
            },
            {
                "name": "언어 사용",
                "description": "영어 표현의 정확성과 유창성",
                "excellent": "정확하고 다양한 표현 사용",
                "good": "대체로 정확한 표현 사용",
                "satisfactory": "기본적인 표현 사용",
                "needs_improvement": "표현이 부정확하거나 제한적"
            },
            {
                "name": "참여도",
                "description": "토론 참여의 적극성과 협력적 태도",
                "excellent": "적극적 참여와 건설적 상호작용",
                "good": "활발한 참여와 협력적 태도",
                "satisfactory": "기본적인 참여",
                "needs_improvement": "소극적 참여 또는 비협력적 태도"
            }
        ]
        
        for criteria in criteria_areas:
            # 영역 제목
            area_title = self.document.add_paragraph()
            area_title.style = 'Section Title'
            area_title.add_run(criteria["name"])
            
            # 설명
            desc_para = self.document.add_paragraph(criteria["description"], style='Custom Body')
            desc_para.runs[0].font.italic = True
            
            # 평가 기준 테이블
            criteria_table = self.document.add_table(rows=5, cols=2)
            criteria_table.style = 'Table Grid'
            
            # 헤더
            criteria_table.cell(0, 0).text = "수준"
            criteria_table.cell(0, 1).text = "기준"
            
            # 평가 수준
            levels = [
                ("우수", criteria["excellent"]),
                ("양호", criteria["good"]),
                ("보통", criteria["satisfactory"]),
                ("개선 필요", criteria["needs_improvement"])
            ]
            
            for i, (level, description) in enumerate(levels, 1):
                criteria_table.cell(i, 0).text = level
                criteria_table.cell(i, 1).text = description
                
                # 헤더 스타일링
                header_cell = criteria_table.cell(i, 0)
                header_para = header_cell.paragraphs[0]
                header_run = header_para.runs[0] if header_para.runs else header_para.add_run(level)
                header_run.font.bold = True
            
            self.document.add_paragraph()  # 공백

# 사용 예시 및 테스트 함수
def test_docx_generator():
    """DOCX 생성기 테스트"""
    
    # 테스트용 설정
    settings = DocumentSettings(
        title="파스칼 영어 토론 교재",
        subtitle="Charlotte's Web 기반 토론 학습",
        author="파스칼 교육팀",
        institution="파스칼 교육원",
        level=EducationLevel.PREPARATION,
        book_title="Charlotte's Web",
        book_author="E.B. White",
        ar_level=4.4,
        creation_date=datetime.now()
    )
    
    # CSV 파일 경로
    csv_path = "/home/ubuntu/pascal_system/Charlotte's_Web_analysis.csv"
    output_path = "/home/ubuntu/pascal_system/Charlotte's_Web_Textbook.docx"
    
    # DOCX 생성
    generator = DOCXGenerator()
    
    try:
        result_path = generator.generate_textbook(csv_path, output_path, settings)
        print(f"Textbook generated successfully: {result_path}")
        return result_path
    except Exception as e:
        print(f"Error generating textbook: {e}")
        return None

if __name__ == "__main__":
    test_docx_generator()
